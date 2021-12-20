import os, sys, getopt
import codecs
from docx import Document

doc = Document()
errorlist = []


def convert(dir, mode='flat', title=None, include=None, exclude=None, encoding='utf-8'):
    print('copy from diretory: ' + dir)

    if title is not None:
        doc.add_heading(title, 1)

    if include is not None:
        inc = include.split('|')
    else:
        inc = None

    if exclude is not None:
        exc = exclude.split('|')
    else:
        exc = None

    if mode == 'flat':
        walkflat(dir, inc, exc, encoding)
    elif mode == 'deep':
        walkdeep(dir, 2, inc, exc, encoding)
    else:
        print('mode is invaild')


def walkflat(dir, inc, exc, encoding):
    currentdir = ''
    for root, dirs, files in os.walk(dir, False):
        for file in files:
            if file == 'pom.xml':
                print(1)
            if (inc is None or os.path.splitext(file)[1][1:] in inc) and (
                    exc is None or os.path.splitext(file)[1][1:] not in exc):
                filepath = os.path.join(root, file).replace('\\', '/')
                try:
                    with codecs.open(filepath, encoding=encoding) as f:
                        content = f.read()
                        thisdir = filepath[len(dir) + 1:filepath.rfind('/')]
                        if currentdir != thisdir:
                            currentdir = thisdir
                            doc.add_heading(thisdir, 2)
                            print('into directory ' + thisdir)
                        doc.add_heading(filepath[filepath.rfind('/') + 1:], 3)
                        doc.add_paragraph(content)
                        doc.add_page_break()
                        print('copied ' + filepath[filepath.rfind('/') + 1:])
                except Exception as e:
                    errorlist.append(filepath)
                    print('read ' + filepath + ' error')
                    print(str(e))


def walkdeep(root, level, inc, exc, encoding):
    for file in os.listdir(root):
        filepath = os.path.join(root, file).replace('\\', '/')
        if os.path.isfile(filepath):
            if (inc is None or os.path.splitext(file)[1][1:] in inc) and (
                    exc is None or os.path.splitext(file)[1][1:] not in exc):
                try:
                    with codecs.open(filepath, encoding=encoding) as f:
                        content = f.read()
                        doc.add_heading(filepath[filepath.rfind('/') + 1:], level)
                        doc.add_paragraph(content)
                        doc.add_page_break()
                        print('copied ' + filepath[filepath.rfind('/') + 1:])
                except Exception as e:
                    errorlist.append(filepath)
                    print('read ' + filepath + ' error')
                    print(str(e))

        else:
            if level<9:
                doc.add_heading(file, level)
                print('into directory ' + file)
                walkdeep(filepath, level + 1, inc, exc, encoding)

if __name__ == '__main__':

  src = None
  mode = 'flat'
  target = None
  include = None
  exclude = None
  encoding = 'utf-8'
  myhelp = 'run.py -s <source directory path> -m \'flat|deep\' -t <target docx file path>\
  -i <include extension of scanned files> -e <exclude extension of scanned files>\
  -c <encoding of the files>'
  argv = sys.argv[1:]

  try:
    opts, args = getopt.getopt(argv,'hs:m:t:i:e:c:',['source=','mode=','target=','include=','exclude=','encoding='])
  except expression as identifier:
    print(myhelp)
    sys.exit(2)

  for opt, arg in opts:
    if opt == '-h':
      print(myhelp)
      sys.exit(2)
    elif opt in ('-s','--source'):
      src = arg
    elif opt in ('-m','--mode'):
      mode = arg
    elif opt in ('-t','--target'):
      target = arg
    elif opt in ('-i','--include'):
      include = arg
    elif opt in ('-e','--exclude'):
      exclude = arg
    elif opt in ('-c','--encoding'):
      encoding = arg

  if src is None or target is None:
    print('source and target is needed')
    sys.exit(2)

  pos = src.find('*')
  if pos == -1:
    convert(src, mode=mode, include=include, exclude=exclude, encoding=encoding)
  else:
    presrc = src[0:pos]
    dirs = os.listdir(presrc)
    for dir in dirs:
      convert(presrc + dir + src[pos+1:], mode=mode, title=dir, include=include, exclude=exclude, encoding=encoding)

  doc.save(target)
  print('\nfinish copying, your document is saved into \"'+target+'\" , thanks for your using!')
  if len(errorlist) != 0:
    print('\nerror file list:\n')
    for e in errorlist:
      print(e)